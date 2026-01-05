"""DOCX fallback extraction using python-docx."""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    parser: str
    error: Optional[str]


def parse_docx(file_path: Path) -> ParsedDocument:
    file_path = Path(file_path)
    try:
        document = Document(file_path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if not text.strip():
            raise ValueError("DOCX extraction returned no text")
        return ParsedDocument(text=text, parser="python-docx", error=None)
    except Exception as exc:  # noqa: BLE001
        logger.warning("DOCX fallback failed for %s: %s", file_path, exc)
        return ParsedDocument(text="", parser="python-docx", error=str(exc))